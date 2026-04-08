import customtkinter as ctk
from tkinter import filedialog, messagebox
import json
from datetime import datetime
import os
import sys
import shutil
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- ФУНКЦИЯ ДЛЯ РАБОТЫ С ПУТЯМИ PYINSTALLER ---
def resource_path(relative_path):
    """ Получает абсолютный путь к ресурсам, работает для dev и для PyInstaller """
    try:
        # PyInstaller создает временную папку _MEIPASS
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# --- ВСПОМОГАТЕЛЬНЫЙ КЛАСС ДЛЯ ПОДСКАЗОК ---
class Tooltip:
    def __init__(self, widget, text):
        self.widget, self.text, self.tip_window = widget, text, None
        self.widget.bind("<Enter>", self.show_tip)
        self.widget.bind("<Leave>", self.hide_tip)

    def show_tip(self, event=None):
        if self.tip_window or not self.text: return
        x, y = self.widget.winfo_rootx() + 20, self.widget.winfo_rooty() + 20
        self.tip_window = tw = ctk.CTkToplevel(self.widget)
        tw.wm_overrideredirect(True)
        tw.wm_geometry(f"+{x}+{y}")
        ctk.CTkLabel(tw, text=self.text, justify="left", fg_color=("#333333", "#F2F2F2"),
                     text_color=("#F2F2F2", "#333333"), corner_radius=6, padx=10, pady=5).pack()

    def hide_tip(self, event=None):
        if self.tip_window: self.tip_window.destroy(); self.tip_window = None

# --- ОКНО РЕДАКТИРОВАНИЯ ---
class EditWindow(ctk.CTkToplevel):
    def __init__(self, master, data, save_callback, delete_callback, is_new=False):
        super().__init__(master)
        self.title("Редактирование" if not is_new else "Новая карточка")
        self.geometry("650x950")
        self.data, self.save_callback, self.delete_callback, self.is_new = data, save_callback, delete_callback, is_new
        self.old_address = self.data["location"].get("raw_address", "")
        
        self.container = ctk.CTkScrollableFrame(self)
        self.container.pack(fill="both", expand=True, padx=10, pady=10)
        self.render_fields()
        self.grab_set()

    def add_label(self, text):
        ctk.CTkLabel(self.container, text=text, font=ctk.CTkFont(size=12, weight="bold")).pack(anchor="w", pady=(10,0))

    def render_fields(self):
        self.add_label("ФИО Резидента:")
        self.name_ent = ctk.CTkEntry(self.container, width=500); self.name_ent.insert(0, self.data.get("resident_name", "")); self.name_ent.pack(pady=5, anchor="w")

        row_stat = ctk.CTkFrame(self.container, fg_color="transparent"); row_stat.pack(fill="x", pady=5)
        ctk.CTkLabel(row_stat, text="Кол-во детей:").pack(side="left", padx=5)
        self.child_ent = ctk.CTkEntry(row_stat, width=60); self.child_ent.insert(0, str(self.data.get("children_count", 0))); self.child_ent.pack(side="left")
        self.village_var = ctk.BooleanVar(value=self.data["location"].get("is_village", False))
        ctk.CTkCheckBox(row_stat, text="Деревня / Село", variable=self.village_var).pack(side="left", padx=20)

        self.add_label("Адрес (Текущий):")
        self.addr_ent = ctk.CTkEntry(self.container, width=500); self.addr_ent.insert(0, self.data["location"].get("raw_address", "")); self.addr_ent.pack(pady=5, anchor="w")

        self.add_label("История переездов:")
        self.loc_hist_txt = ctk.CTkTextbox(self.container, height=80, width=500)
        self.loc_hist_txt.pack(pady=5, anchor="w")
        self.loc_hist_txt.insert("1.0", "\n".join(self.data.get("location_history", [])))
        
        self.add_label("Оборудование (Модель / Кол-во / Дата установки ДД.ММ.ГГГГ):")
        dev_f = ctk.CTkFrame(self.container, fg_color="transparent"); dev_f.pack(fill="x")
        self.mod_ent = ctk.CTkEntry(dev_f, width=200); self.mod_ent.insert(0, self.data["device"]["model"]); self.mod_ent.pack(side="left", padx=2)
        self.cnt_ent = ctk.CTkEntry(dev_f, width=50); self.cnt_ent.insert(0, str(self.data["device"]["count"])); self.cnt_ent.pack(side="left", padx=2)
        self.inst_ent = ctk.CTkEntry(dev_f, width=120); self.inst_ent.insert(0, self.data["device"]["install_date"]); self.inst_ent.pack(side="left", padx=2)

        self.add_label("Даты проверок (через запятую):")
        hist_dates = ", ".join([h['date'] for h in self.data["check_history"]])
        self.checks_ent = ctk.CTkEntry(self.container, width=500); self.checks_ent.insert(0, hist_dates); self.checks_ent.pack(pady=5, anchor="w")

        self.add_label("Социальные заметки:")
        self.soc_ent = ctk.CTkEntry(self.container, width=500); self.soc_ent.insert(0, self.data["notes"].get("social", "")); self.soc_ent.pack(pady=5, anchor="w")
        self.add_label("Технические заметки:")
        self.tech_ent = ctk.CTkEntry(self.container, width=500); self.tech_ent.insert(0, self.data["notes"].get("tech", "")); self.tech_ent.pack(pady=5, anchor="w")

        if not self.is_new:
            actions_f = ctk.CTkFrame(self.container, fg_color="transparent"); actions_f.pack(fill="x", pady=20)
            btn_txt = "В АРХИВ" if self.data.get("status") == "active" else "ВЕРНУТЬ В АКТИВ"
            ctk.CTkButton(actions_f, text=btn_txt, width=240, command=self.toggle_status).pack(side="left", padx=5)
            ctk.CTkButton(actions_f, text="УДАЛИТЬ", width=240, fg_color="#942d2d", command=self.confirm_delete).pack(side="right", padx=5)

        ctk.CTkButton(self, text="СОХРАНИТЬ", fg_color="#1f538d", height=50, command=self.save).pack(fill="x", padx=20, pady=15)

    def toggle_status(self):
        self.data["status"] = "archived" if self.data.get("status") == "active" else "active"
        messagebox.showinfo("Статус", f"Обновлено: {self.data['status']}")

    def confirm_delete(self):
        if messagebox.askyesno("Удаление", "Удалить эту карточку навсегда?"):
            self.delete_callback(self.data["id"]); self.destroy()

    def save(self):
        new_addr = self.addr_ent.get()
        if not self.is_new and self.old_address and new_addr != self.old_address:
            if self.old_address not in self.data.get("location_history", []):
                self.data.setdefault("location_history", []).append(self.old_address)

        self.data["resident_name"] = self.name_ent.get()
        try: self.data["children_count"] = int(self.child_ent.get() or 0)
        except: self.data["children_count"] = 0
        self.data["location"]["raw_address"] = new_addr
        self.data["location"]["is_village"] = self.village_var.get()
        self.data["location_history"] = [l.strip() for l in self.loc_hist_txt.get("1.0", "end-1c").split("\n") if l.strip()]
        self.data["device"]["model"] = self.mod_ent.get()
        try: self.data["device"]["count"] = int(self.cnt_ent.get() or 0)
        except: self.data["device"]["count"] = 0
        self.data["device"]["install_date"] = self.inst_ent.get()
        dates = [d.strip() for d in self.checks_ent.get().split(",") if d.strip()]
        self.data["check_history"] = [{"date": d} for d in dates]
        self.data["notes"] = {"social": self.soc_ent.get(), "tech": self.tech_ent.get()}
        self.save_callback(self.data, self.is_new); self.destroy()

# --- КАРТОЧКА ---
class ResidentCard(ctk.CTkFrame):
    def __init__(self, master, data, app, is_selected=False):
        super().__init__(master)
        self.data, self.app = data, app
        self.configure(fg_color=("#F2F2F2", "#2B2B2B"), corner_radius=10)
        self.is_selected = ctk.BooleanVar(value=is_selected)
        self.render()

    def render(self):
        for child in self.winfo_children(): child.destroy()
        ctk.CTkCheckBox(self, text="", variable=self.is_selected, width=20, command=self.on_toggle).pack(side="left", padx=(15, 5))
        left_f = ctk.CTkFrame(self, fg_color="transparent"); left_f.pack(side="left", fill="both", expand=True, padx=5, pady=10)
        header = ctk.CTkFrame(left_f, fg_color="transparent"); header.pack(fill="x")
        is_act = self.data.get("status") == "active"
        ctk.CTkLabel(header, text="●" if is_act else "○", text_color="#2fa572" if is_act else "#942d2d").pack(side="left")
        ctk.CTkLabel(header, text=f"{self.data['resident_name']} | Детей: {self.data.get('children_count', 0)}", font=ctk.CTkFont(size=14, weight="bold")).pack(side="left", padx=5)

        if self.data["location"].get("is_village"):
            ctk.CTkLabel(header, text="[СЕЛО]", text_color="#c9a100", font=("Arial", 10, "bold")).pack(side="left", padx=5)

        notes = self.data.get("notes", {})
        if notes.get("social") or notes.get("tech"):
            q = ctk.CTkLabel(header, text=" [?]", text_color="#3a7ebf", cursor="hand2"); q.pack(side="left")
            Tooltip(q, f"СОЦ: {notes.get('social','—')}\nТЕХ: {notes.get('tech','—')}")

        ctk.CTkLabel(left_f, text=self.data["location"]["raw_address"], font=("Arial", 11), text_color="gray", anchor="w").pack(fill="x")

        info_f = ctk.CTkFrame(self, fg_color="transparent"); info_f.pack(side="left", padx=20)
        ctk.CTkLabel(info_f, text=f"{self.data['device'].get('count',0)} шт.", font=("Arial", 18, "bold"), text_color="#3a7ebf").pack()
        ctk.CTkLabel(info_f, text="ДАТЧИКОВ", font=("Arial", 8)).pack()

        check_f = ctk.CTkFrame(self, fg_color="transparent"); check_f.pack(side="left", padx=20)
        last_date = self.data["check_history"][-1]["date"] if self.data.get("check_history") else "—"
        ctk.CTkLabel(check_f, text=last_date, font=("Arial", 14, "bold")).pack()
        ctk.CTkLabel(check_f, text="ПРОВЕРКА", font=("Arial", 8)).pack()

        ctk.CTkButton(self, text="ИЗМЕНИТЬ", width=80, command=lambda: EditWindow(self.app, self.data, self.app.on_save, self.app.on_delete)).pack(side="right", padx=15)

    def on_toggle(self):
        if self.is_selected.get(): self.app.selected_ids.add(self.data["id"])
        else: self.app.selected_ids.discard(self.data["id"])
        self.app.update_selection_counter()

# --- ГЛАВНОЕ ПРИЛОЖЕНИЕ ---
class App(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("АДПИ Менеджер")
        self.geometry("1450x950")
        
        # Для реестра используем путь рядом с исполняемым файлом, а не во временной папке
        if getattr(sys, 'frozen', False):
            base_dir = os.path.dirname(sys.executable)
        else:
            base_dir = os.path.dirname(os.path.abspath(__file__))
            
        self.reg_dir = os.path.join(base_dir, "registry")
        if not os.path.exists(self.reg_dir): os.makedirs(self.reg_dir)

        self.full_data, self._pool, self.selected_ids = [], [], set()
        self.file_path, self.shown, self.limit = None, 0, 25
        self.f_status, self.f_loc_type = "all", "Все"
        self.f_overdue = ctk.BooleanVar()

        self.setup_ui()
        self.auto_load_registry()
        self.monitor_scroll()

    def setup_ui(self):
        top = ctk.CTkFrame(self); top.pack(fill="x", padx=20, pady=10)
        
        r1 = ctk.CTkFrame(top, fg_color="transparent"); r1.pack(fill="x", padx=10, pady=5)
        self.search = ctk.CTkEntry(r1, placeholder_text="Поиск...", width=400)
        self.search.pack(side="left", padx=5); self.search.bind("<KeyRelease>", lambda e: self.apply_filters())
        ctk.CTkButton(r1, text="+ НОВАЯ СЕМЬЯ", fg_color="#2fa572", command=self.add_new).pack(side="left", padx=5)
        ctk.CTkButton(r1, text="ЗАГРУЗИТЬ JSON", command=self.load_file_dialog).pack(side="right", padx=5)
        ctk.CTkButton(r1, text="EXCEL ЭКСПОРТ", fg_color="#1f538d", command=self.export_excel_menu).pack(side="right", padx=5)

        r2 = ctk.CTkFrame(top, fg_color="transparent"); r2.pack(fill="x", padx=10, pady=5)
        self.seg_stat = ctk.CTkSegmentedButton(r2, values=["Все", "Активные", "Архив"], command=self.set_stat_f); self.seg_stat.set("Все"); self.seg_stat.pack(side="left", padx=5)
        self.seg_loc = ctk.CTkSegmentedButton(r2, values=["Все", "Город", "Деревня"], command=self.set_loc_f); self.seg_loc.set("Все"); self.seg_loc.pack(side="left", padx=10)
        
        ctk.CTkLabel(r2, text="Детей от:").pack(side="left", padx=2)
        self.f_child_min = ctk.CTkEntry(r2, width=40); self.f_child_min.pack(side="left")
        ctk.CTkLabel(r2, text="Датчиков от:").pack(side="left", padx=(10, 2))
        self.f_min_cnt = ctk.CTkEntry(r2, width=40); self.f_min_cnt.pack(side="left")
        ctk.CTkCheckBox(r2, text="Просрочено", variable=self.f_overdue, command=self.apply_filters).pack(side="left", padx=15)

        r3 = ctk.CTkFrame(top, fg_color="transparent"); r3.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(r3, text="Установка с:").pack(side="left", padx=5)
        self.f_date_start = ctk.CTkEntry(r3, placeholder_text="ДД.ММ.ГГГГ", width=100); self.f_date_start.pack(side="left")
        ctk.CTkLabel(r3, text="по:").pack(side="left", padx=5)
        self.f_date_end = ctk.CTkEntry(r3, placeholder_text="ДД.ММ.ГГГГ", width=100); self.f_date_end.pack(side="left")
        ctk.CTkButton(r3, text="ПРИМЕНИТЬ", width=120, command=self.apply_filters).pack(side="left", padx=20)
        ctk.CTkButton(r3, text="ПЕЧАТЬ СПИСКА (WORD)", fg_color="#c9a100", text_color="black", command=self.export_trip_list).pack(side="right", padx=5)

        sel_p = ctk.CTkFrame(top, fg_color=("#DBDBDB", "#2B2B2B"), corner_radius=8); sel_p.pack(fill="x", padx=10, pady=10)
        self.lbl_select = ctk.CTkLabel(sel_p, text="ВЫБРАНО: 0", font=ctk.CTkFont(weight="bold")); self.lbl_select.pack(side="left", padx=20, pady=5)
        ctk.CTkButton(sel_p, text="СБРОСИТЬ ВЫБОР", fg_color="#942d2d", width=120, command=self.clear_selection).pack(side="right", padx=5)

        self.scroll = ctk.CTkScrollableFrame(self, label_text="Реестр семей"); self.scroll.pack(fill="both", expand=True, padx=20, pady=10)

    def set_stat_f(self, v): self.f_status = {"Все":"all", "Активные":"active", "Архив":"archived"}[v]; self.apply_filters()
    def set_loc_f(self, v): self.f_loc_type = v; self.apply_filters()

    def apply_filters(self, *args):
        q, res = self.search.get().lower(), self.full_data
        if q: res = [i for i in res if q in i["resident_name"].lower() or q in i["location"]["raw_address"].lower()]
        if self.f_status != "all": res = [i for i in res if i.get("status") == self.f_status]
        if self.f_loc_type != "Все":
            is_v = self.f_loc_type == "Деревня"; res = [i for i in res if i["location"].get("is_village") == is_v]

        try:
            c_min = int(self.f_child_min.get()) if self.f_child_min.get() else 0
            res = [i for i in res if i.get("children_count", 0) >= c_min]
            d_min = int(self.f_min_cnt.get()) if self.f_min_cnt.get() else 0
            res = [i for i in res if i["device"].get("count", 0) >= d_min]
        except: pass

        ds_str, de_str = self.f_date_start.get(), self.f_date_end.get()
        if ds_str or de_str:
            try:
                ds = datetime.strptime(ds_str, "%d.%m.%Y") if ds_str else datetime.min
                de = datetime.strptime(de_str, "%d.%m.%Y") if de_str else datetime.max
                res = [i for i in res if i["device"].get("install_date") and ds <= datetime.strptime(i["device"]["install_date"], "%d.%m.%Y") <= de]
            except: pass

        if self.f_overdue.get():
            now = datetime.now()
            res = [i for i in res if not i.get("check_history") or (now - datetime.strptime(i["check_history"][-1]["date"], "%d.%m.%Y")).days > 365]

        res.sort(key=lambda x: datetime.strptime(x["check_history"][-1]["date"], "%d.%m.%Y").timestamp() if x.get("check_history") else 0)
        self._pool = res; self.refresh_list()

    def export_trip_list(self):
        if not self.selected_ids: 
            messagebox.showwarning("!", "Выберите семьи галочками!")
            return
            
        path = filedialog.asksaveasfilename(
            defaultextension=".docx", 
            initialfile=f"Установка_АДПИ_{datetime.now().strftime('%d_%m_%Y')}.docx"
        )
        if not path: return
        
        doc = Document()
        
        # Настройка шрифта по умолчанию
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(11)
        
        # Заголовок по центру
        title_text = f"Установка АДПИ на {datetime.now().strftime('%d.%m.%Y')} год"
        title = doc.add_paragraph(title_text)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.runs[0]
        run.bold = True
        run.font.size = Pt(14)
        
        # Создание таблицы (8 колонок)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        
        # Шапка таблицы
        hdr_cells = table.rows[0].cells
        headers = [
            "№", 
            "ФИО родителя", 
            "Адрес проживания", 
            "Номер телефона", 
            "АДПИ (шт)", 
            "Дети (кол-во)", 
            "ФИО детей, дата рождения", 
            "О жилищных условиях"
        ]
        
        for idx, text in enumerate(headers):
            hdr_cells[idx].text = text
            paragraph = hdr_cells[idx].paragraphs[0]
            if paragraph.runs:
                run = paragraph.runs[0]
                run.bold = True
                run.font.size = Pt(10)

        # Получаем данные выбранных семей
        selected_items = [i for i in self.full_data if i["id"] in self.selected_ids]
        
        for item in selected_items:
            row_cells = table.add_row().cells
            
            # Оставляем ПУСТЫМИ согласно требованию:
            row_cells[0].text = "" # №
            
            # Заполняем только базовую информацию для идентификации:
            row_cells[1].text = str(item.get('resident_name', ''))
            row_cells[2].text = str(item['location'].get('raw_address', ''))
            
            # Оставляем ПУСТЫМИ для ручного заполнения (включая АДПИ и Дети):
            row_cells[3].text = "" # Номер телефона
            row_cells[4].text = "" # АДПИ (шт) - ТЕПЕРЬ ПУСТО
            row_cells[5].text = "" # Дети (кол-во) - ТЕПЕРЬ ПУСТО
            row_cells[6].text = "" # ФИО детей
            row_cells[7].text = "" # О жилищных условиях
            
            # Применяем шрифт ко всей строке
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)

        try:
            doc.save(path)
            messagebox.showinfo("OK", "Документ Word сформирован.")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл (возможно, он открыт в Word): {e}")

    def export_excel_menu(self):
        p = filedialog.asksaveasfilename(defaultextension=".xlsx")
        if not p: return
        target_list = [i for i in self.full_data if i["id"] in self.selected_ids] if self.selected_ids else self._pool
        rows = []
        for i in target_list:
            rows.append({
                "ФИО": i.get("resident_name"), "Детей": i.get("children_count"), "Адрес": i["location"].get("raw_address"),
                "История адресов": ", ".join(i.get("location_history", [])), "Кол-во АДПИ": i["device"].get("count"),
                "Дата установки": i["device"].get("install_date"), "Последняя проверка": i["check_history"][-1]["date"] if i.get("check_history") else "Нет"
            })
        pd.DataFrame(rows).to_excel(p, index=False); messagebox.showinfo("OK", "Экспорт завершен")

    def refresh_list(self): 
        for c in self.scroll.winfo_children(): c.destroy()
        try: self.scroll._parent_canvas.yview_moveto(0)
        except: pass
        self.shown = 0; self.load_batch()

    def load_batch(self):
        start, end = self.shown, min(self.shown + self.limit, len(self._pool))
        for i in range(start, end):
            ResidentCard(self.scroll, self._pool[i], self, self._pool[i]["id"] in self.selected_ids).pack(fill="x", padx=10, pady=5)
        self.shown = end

    def update_selection_counter(self): self.lbl_select.configure(text=f"ВЫБРАНО: {len(self.selected_ids)}")
    def clear_selection(self): self.selected_ids.clear(); self.update_selection_counter(); self.refresh_list()
    
    def on_save(self, data, is_new):
        if is_new: self.full_data.append(data)
        with open(self.file_path, 'w', encoding='utf-8') as f: json.dump(self.full_data, f, ensure_ascii=False, indent=4)
        self.apply_filters()

    def on_delete(self, rid):
        self.full_data = [i for i in self.full_data if i["id"] != rid]
        self.selected_ids.discard(rid); self.on_save(None, False)

    def load_data(self, path):
        self.file_path = path
        try:
            with open(path, 'r', encoding='utf-8') as f: self.full_data = json.load(f)
            self.apply_filters()
        except: pass

    def auto_load_registry(self):
        files = [f for f in os.listdir(self.reg_dir) if f.endswith(".json")]
        if files: 
            newest_file = max([os.path.join(self.reg_dir, f) for f in files], key=os.path.getmtime)
            self.load_data(newest_file)

    def load_file_dialog(self):
        p = filedialog.askopenfilename(filetypes=[("JSON", "*.json")])
        if p:
            dest = os.path.join(self.reg_dir, os.path.basename(p))
            if os.path.abspath(p) != os.path.abspath(dest): shutil.copy(p, dest)
            self.load_data(dest)

    def add_new(self):
        nid = max([i["id"] for i in self.full_data]) + 1 if self.full_data else 1
        new_i = {"id": nid, "resident_name": "", "children_count": 0, "location": {"raw_address": "", "is_village": False}, "location_history": [], "device": {"model": "", "count": 0, "install_date": ""}, "check_history": [], "status": "active", "notes": {"social": "", "tech": ""}}
        EditWindow(self, new_i, self.on_save, self.on_delete, True)

    def monitor_scroll(self):
        try:
            if self.scroll._scrollbar.get()[1] > 0.9 and self.shown < len(self._pool): self.load_batch()
        except: pass
        self.after(300, self.monitor_scroll)

if __name__ == "__main__":
    App().mainloop()